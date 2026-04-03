import io
import os
import sqlite3
from datetime import date, datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt

st.set_page_config(page_title="APP Área Comercial Buses y Vans", layout="wide")

TEMPLATE_FILE = "Propuesta_FOTON_Electrico_U9.docx"
DB_FILE = "cotizaciones.db"

COTIZANTES = {
    "Diego Vejar": {
        "prefijo": "DV",
        "nombre_firma": "Diego Vejar",
        "cargo": "Subgerente Comercial Buses y Vans",
        "correo": "dvejar@andesmotor.cl",
        "telefono_fijo": "",
        "telefono_movil": "981774604",
    },
    "Sergio Silva": {
        "prefijo": "SS",
        "nombre_firma": "Sergio Silva",
        "cargo": "Ejecutivo de ventas Zona Sur Buses y Vans",
        "correo": "sergio.silva@andesmotor.cl",
        "telefono_fijo": "",
        "telefono_movil": "",
    },
    "Alvaro Correa": {
        "prefijo": "AC",
        "nombre_firma": "Alvaro Correa",
        "cargo": "Ejecutivo de ventas Zona Norte Buses y Vans",
        "correo": "alvaro.correa@andesmotor.cl",
        "telefono_fijo": "",
        "telefono_movil": "",
    },
    "Fabian Orellana": {
        "prefijo": "FO",
        "nombre_firma": "Fabian Orellana",
        "cargo": "Product Manager Senior Zona Buses y Vans",
        "correo": "forellana@andesmotor.cl",
        "telefono_fijo": "",
        "telefono_movil": "989356449",
    },
    "Rodrigo Sepulveda": {
        "prefijo": "RS",
        "nombre_firma": "Rodrigo Sepulveda",
        "cargo": "Gerente de Buses y Vans",
        "correo": "rsepulveda@andesmotor.cl",
        "telefono_fijo": "",
        "telefono_movil": "979783254",
    },
}


# =========================================================
# BASE DE DATOS
# =========================================================
def get_connection():
    return sqlite3.connect(DB_FILE, check_same_thread=False)


def init_db():
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS cotizaciones (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            fecha TEXT NOT NULL,
            cliente TEXT NOT NULL,
            cotizante TEXT NOT NULL,
            prefijo TEXT NOT NULL,
            correlativo INTEGER NOT NULL,
            numero_cotizacion TEXT NOT NULL UNIQUE,
            cantidad_unidades INTEGER NOT NULL,
            precio_unitario REAL NOT NULL,
            total_negocio REAL NOT NULL,
            contrato_mantto TEXT,
            texto_mantto TEXT,
            creado_en TEXT NOT NULL
        )
    """)

    conn.commit()
    conn.close()


def obtener_siguiente_correlativo(cotizante: str) -> int:
    prefijo = COTIZANTES[cotizante]["prefijo"]
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT MAX(correlativo)
        FROM cotizaciones
        WHERE prefijo = ?
    """, (prefijo,))

    row = cur.fetchone()
    conn.close()

    ultimo = row[0] if row and row[0] is not None else 0
    return ultimo + 1


def guardar_cotizacion(datos: dict):
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        INSERT INTO cotizaciones (
            fecha, cliente, cotizante, prefijo, correlativo, numero_cotizacion,
            cantidad_unidades, precio_unitario, total_negocio,
            contrato_mantto, texto_mantto, creado_en
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        datos["fecha_iso"],
        datos["cliente"],
        datos["cotizante"],
        datos["prefijo"],
        datos["correlativo"],
        datos["numero_cotizacion"],
        datos["cantidad_unidades"],
        datos["precio_unitario_raw"],
        datos["total_negocio_raw"],
        datos["contrato_mantto"],
        datos["texto_mantto"],
        datetime.now().isoformat(timespec="seconds")
    ))

    conn.commit()
    conn.close()


def cargar_historial():
    conn = get_connection()
    df = pd.read_sql_query("""
        SELECT
            fecha,
            cliente,
            cotizante,
            numero_cotizacion,
            cantidad_unidades,
            precio_unitario,
            total_negocio,
            creado_en
        FROM cotizaciones
        ORDER BY id DESC
    """, conn)
    conn.close()
    return df


# =========================================================
# UTILIDADES WORD
# =========================================================
def fecha_larga_es(fecha: date) -> str:
    dias = ["lunes", "martes", "miércoles", "jueves", "viernes", "sábado", "domingo"]
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
    ]
    return f"Santiago, {dias[fecha.weekday()]} {fecha.day} de {meses[fecha.month - 1]} de {fecha.year}"


def usd_fmt(valor: float) -> str:
    return f"USD$ {valor:,.0f}".replace(",", ".")


def limpiar_parrafo(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def escribir_parrafo_simple(paragraph, texto, bold=False, font_size=Pt(11)):
    limpiar_parrafo(paragraph)
    run = paragraph.add_run(texto)
    run.bold = bold
    run.font.size = font_size


def escribir_parrafo_precio(paragraph, precio_texto):
    limpiar_parrafo(paragraph)

    r1 = paragraph.add_run("•    Valor unitario por bus: ")
    r1.font.size = Pt(11)

    r2 = paragraph.add_run(precio_texto)
    r2.bold = True
    r2.font.size = Pt(11)

    r3 = paragraph.add_run(" + IVA.")
    r3.font.size = Pt(11)


def insertar_parrafo_despues(paragraph, texto):
    nuevo = paragraph.insert_paragraph_before("")
    escribir_parrafo_simple(nuevo, texto)
    return nuevo


def reemplazar_en_parrafos(doc, datos):
    cotizante_info = COTIZANTES[datos["cotizante"]]

    for p in doc.paragraphs:
        txt = p.text.strip()

        if txt.startswith("Santiago,"):
            escribir_parrafo_simple(p, datos["fecha_larga"])

        elif "Cotización n°" in txt:
            escribir_parrafo_simple(p, f'Cotización n° {datos["numero_cotizacion"]}')

        elif txt == "Conecta":
            escribir_parrafo_simple(p, datos["cliente"])

        elif txt.startswith("1) DESCRIPCIÓN:"):
            nuevo = (
                f'1) DESCRIPCIÓN: {datos["cantidad_unidades"]} Buses Eléctricos, '
                f'Marca FOTON, Modelo U9 de carrocería monocasco Urbano estándar RED, año 2026.'
            )
            escribir_parrafo_simple(p, nuevo)

        elif "Valor unitario por bus:" in txt:
            escribir_parrafo_precio(p, datos["precio_unitario"])

        elif txt.startswith("I.") or txt.startswith("II.") or txt.startswith("III."):
            escribir_parrafo_simple(p, "")

        elif "Rodrigo Sepúlveda Toepfer" in txt or "Rodrigo Sepulveda Toepfer" in txt:
            escribir_parrafo_simple(p, cotizante_info["nombre_firma"])

        elif "Gerente de Buses y Vans" in txt or "Gerente de Buses y Vans (Iveco)" in txt:
            escribir_parrafo_simple(p, cotizante_info["cargo"])

        elif "@andesmotor.cl" in txt:
            escribir_parrafo_simple(p, cotizante_info["correo"])

        elif txt.startswith("T:"):
            valor = f'T: {cotizante_info["telefono_fijo"]}' if cotizante_info["telefono_fijo"] else ""
            escribir_parrafo_simple(p, valor)

        elif txt.startswith("M:"):
            valor = f'M: {cotizante_info["telefono_movil"]}' if cotizante_info["telefono_movil"] else ""
            escribir_parrafo_simple(p, valor)


def reemplazar_en_tablas(doc, datos):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip() == "Conecta":
                        escribir_parrafo_simple(p, datos["cliente"])


def insertar_bloque_mantenimiento(doc, texto_mantto):
    parrafo_precio = None

    for p in doc.paragraphs:
        if "Valor unitario por bus:" in p.text:
            parrafo_precio = p
            break

    if parrafo_precio is None:
        return

    lineas = [x.strip() for x in texto_mantto.split("\n") if x.strip()]
    ref = parrafo_precio

    for linea in reversed(lineas):
        nuevo = ref.insert_paragraph_before("")
        escribir_parrafo_simple(nuevo, linea)


def generar_docx(datos):
    if not os.path.exists(TEMPLATE_FILE):
        raise FileNotFoundError(f"No se encontró la plantilla: {TEMPLATE_FILE}")

    doc = Document(TEMPLATE_FILE)

    reemplazar_en_parrafos(doc, datos)
    reemplazar_en_tablas(doc, datos)
    insertar_bloque_mantenimiento(doc, datos["texto_mantto"])

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# =========================================================
# APP
# =========================================================
init_db()

st.title("APP Área Comercial Buses y Vans")
st.subheader("Generador de Cotización FOTON U9")

tab1, tab2 = st.tabs(["Nueva cotización", "Historial"])

with tab1:
    with st.form("form_cotizacion"):
        c1, c2 = st.columns(2)

        with c1:
            fecha = st.date_input("Fecha", value=date.today())
            cliente = st.text_input("Cliente", value="")
            cotizante = st.selectbox("Cotizante", list(COTIZANTES.keys()))

            correlativo_sugerido = obtener_siguiente_correlativo(cotizante)
            numero_sugerido = f'{COTIZANTES[cotizante]["prefijo"]}-{correlativo_sugerido:02d}'

            st.info(f"Próximo correlativo automático: {numero_sugerido}")

        with c2:
            cantidad_unidades = st.number_input("Cantidad de unidades", min_value=1, value=1, step=1)
            precio_unitario = st.number_input("Precio unitario USD", min_value=0.0, value=130491.0, step=1000.0)
            contrato_mantto = st.text_input("Contrato mantto", value="48 meses")

        texto_mantto = st.text_area(
            "Bloque mantenimiento / beneficios",
            value=(
                "I. La oferta incluye 48 meses de mantenimiento preventivo y correctivo "
                "(Correctivo de desgaste Disco y Pastilla de frenos) sin costo para el cliente, "
                "con el fin de entregar conocimientos técnicos y aprendizaje continuo del mantenimiento "
                "para este tipo de vehículos, durante este periodo.\n\n"
                "II. Adicionalmente se entregarán USD 1.500.- por bus, para la compra de repuestos de desgaste "
                "a elección del cliente (este ítem deberá ser utilizado dentro de los primeros 12 meses realizada la entrega de flota).\n\n"
                "III. La oferta incluye 8 años de telemetría sin costo para el cliente."
            ),
            height=220
        )

        enviar = st.form_submit_button("Generar y guardar cotización")

    if enviar:
        if not cliente.strip():
            st.error("Debes ingresar el nombre del cliente.")
        else:
            info = COTIZANTES[cotizante]
            correlativo = obtener_siguiente_correlativo(cotizante)
            numero_cotizacion = f'{info["prefijo"]}-{correlativo:02d}'
            total_negocio = precio_unitario * cantidad_unidades

            datos = {
                "fecha_iso": fecha.isoformat(),
                "fecha_larga": fecha_larga_es(fecha),
                "cliente": cliente.strip(),
                "cotizante": cotizante,
                "prefijo": info["prefijo"],
                "correlativo": correlativo,
                "numero_cotizacion": numero_cotizacion,
                "cantidad_unidades": int(cantidad_unidades),
                "precio_unitario": usd_fmt(precio_unitario),
                "precio_unitario_raw": float(precio_unitario),
                "total_negocio_raw": float(total_negocio),
                "contrato_mantto": contrato_mantto,
                "texto_mantto": texto_mantto,
            }

            try:
                archivo = generar_docx(datos)
                guardar_cotizacion(datos)

                st.success(f"Cotización {numero_cotizacion} generada y guardada correctamente.")
                st.write(f"**Cliente:** {cliente.strip()}")
                st.write(f"**Cotizante:** {cotizante}")
                st.write(f"**Precio unitario:** {usd_fmt(precio_unitario)}")
                st.write(f"**Total negocio:** {usd_fmt(total_negocio)}")

                st.download_button(
                    "Descargar cotización Word",
                    data=archivo,
                    file_name=f"Cotizacion_{numero_cotizacion}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.rerun()

            except Exception as e:
                st.error(f"Ocurrió un error al generar la cotización: {e}")

with tab2:
    st.subheader("Historial de cotizaciones")

    try:
        historial = cargar_historial()

        if historial.empty:
            st.info("Aún no hay cotizaciones registradas.")
        else:
            historial["precio_unitario"] = historial["precio_unitario"].apply(usd_fmt)
            historial["total_negocio"] = historial["total_negocio"].apply(usd_fmt)
            st.dataframe(historial, use_container_width=True)

            csv_data = historial.to_csv(index=False).encode("utf-8-sig")
            st.download_button(
                "Descargar historial CSV",
                data=csv_data,
                file_name="historial_cotizaciones.csv",
                mime="text/csv"
            )

    except Exception as e:
        st.error(f"No fue posible cargar el historial: {e}")